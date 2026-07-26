from pathlib import Path

from PIL import Image

from drive_stripper import metadata

_EXIF_ARTIST_TAG = 315


def _make_image_with_exif(path: Path) -> None:
    img = Image.new("RGB", (4, 4), color="red")
    exif = img.getexif()
    exif[_EXIF_ARTIST_TAG] = "Acme Corp"
    img.save(path, exif=exif)


def test_strip_image_metadata_removes_exif(tmp_path: Path):
    source = tmp_path / "photo.jpg"
    dest = tmp_path / "photo.clean.jpg"
    _make_image_with_exif(source)

    with Image.open(source) as img:
        assert len(img.getexif()) > 0

    handled = metadata.strip_metadata(source, dest)
    assert handled is True

    with Image.open(dest) as img:
        assert len(img.getexif()) == 0


def test_strip_docx_metadata_clears_author(tmp_path: Path):
    from docx import Document

    source = tmp_path / "doc.docx"
    dest = tmp_path / "doc.clean.docx"

    doc = Document()
    doc.core_properties.author = "Acme Corp"
    doc.add_paragraph("hello world")
    doc.save(source)

    handled = metadata.strip_metadata(source, dest)
    assert handled is True

    clean_doc = Document(str(dest))
    assert clean_doc.core_properties.author == ""


def test_strip_xlsx_metadata_clears_creator(tmp_path: Path):
    from openpyxl import Workbook, load_workbook

    source = tmp_path / "book.xlsx"
    dest = tmp_path / "book.clean.xlsx"

    wb = Workbook()
    wb.properties.creator = "Acme Corp"
    wb.save(source)

    handled = metadata.strip_metadata(source, dest)
    assert handled is True

    clean_wb = load_workbook(str(dest))
    # openpyxl round-trips a cleared (empty-string) property as None - both mean "no creator"
    assert clean_wb.properties.creator in ("", None)


def test_unsupported_type_is_copied_through(tmp_path: Path):
    source = tmp_path / "notes.txt"
    dest = tmp_path / "notes.clean.txt"
    source.write_text("just plain text")

    handled = metadata.strip_metadata(source, dest)
    assert handled is False
    assert dest.read_text() == "just plain text"
