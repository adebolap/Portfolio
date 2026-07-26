"""Extract and rewrite the visible text content of a file.

Metadata stripping (see ``metadata.py``) only handles what's bolted onto a
file behind the scenes. Proprietary terms and secrets usually live in the
visible text itself, so this module knows how to pull that text out for
scanning, and - for formats where it's straightforward - write redacted or
scaffolded text back in.

PDF is read-only here: rewriting a PDF's content streams in place is a much
bigger undertaking than this tool's scope, so PDF redaction produces a
sanitized ``.txt`` sidecar of the extracted text rather than editing the PDF.
"""

from __future__ import annotations

from pathlib import Path

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".py", ".log", ".yaml", ".yml"}


def is_text_extractable(path: Path) -> bool:
    return path.suffix.lower() in TEXT_SUFFIXES | {".docx", ".pdf", ".xlsx"}


def read_text(path: Path) -> str | None:
    """Best-effort plain-text extraction. Returns None for formats with no text layer."""
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return path.read_text(errors="replace")
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(str(path), data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                lines.append(" ".join(str(c) for c in row if c is not None))
        return "\n".join(lines)
    return None


def write_text(path: Path, text: str) -> None:
    path.write_text(text)


def write_docx_text(source: Path, destination: Path, transform) -> None:
    """Rewrite every paragraph's text in a copy of ``source`` via ``transform(str) -> str``.

    Applies at paragraph granularity: the first run keeps the replacement
    text and any remaining runs in that paragraph are cleared, so exact
    intra-paragraph formatting boundaries are not preserved.
    """
    from docx import Document

    doc = Document(str(source))
    for paragraph in doc.paragraphs:
        if not paragraph.runs:
            continue
        new_text = transform(paragraph.text)
        if new_text == paragraph.text:
            continue
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    doc.save(str(destination))


def write_xlsx_text(source: Path, destination: Path, transform) -> None:
    """Rewrite every string cell in a copy of ``source`` via ``transform(str) -> str``."""
    from openpyxl import load_workbook

    wb = load_workbook(str(source))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value = transform(cell.value)
    wb.save(str(destination))
